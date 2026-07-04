"""Memória-biztos szövegkinyerés .pdf / .docx / .xlsx / .txt fájlokból.

Minden loader generátor: szakaszonként yield-el, sosem materializálja a teljes
dokumentumot egyetlen stringként. A hívó (ingestion) felel a globális
karakterlimit betartásáért.
"""
from __future__ import annotations

import io
from typing import Iterator

from .config import settings


def load_txt(data: bytes) -> Iterator[str]:
    text = data.decode("utf-8", errors="replace")
    # 64k-s szeletekben adjuk tovább, hogy a splitter kezelhető darabokat kapjon
    for i in range(0, len(text), 65536):
        yield text[i : i + 65536]


def load_pdf(data: bytes) -> Iterator[str]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            yield text


def load_docx(data: bytes) -> Iterator[str]:
    import docx

    document = docx.Document(io.BytesIO(data))
    for para in document.paragraphs:
        if para.text.strip():
            yield para.text
    # Táblázatok: sorok tabulátorral összefűzve — a struktúra részben megmarad
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                yield "\t".join(cells)


def load_xlsx(data: bytes) -> Iterator[str]:
    """read_only + iter_rows: az openpyxl soronként streamel, nem tölti be
    a teljes munkafüzetet a memóriába. Cellánkénti karakterplafon védi a
    degenerált (óriás szöveget tartalmazó) cellák elleni esetet.

    Az adatsorokat "Fejléc: érték | Fejléc: érték" párokká alakítjuk — az LLM
    nyers tab-elválasztott soroknál elcsúszik az oszlopokon, a felcímkézett
    pároknál nem.
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        for ws in wb.worksheets:
            yield f"[Munkalap: {ws.title}]"
            headers: list[str] | None = None
            for row in ws.iter_rows(values_only=True):
                cells = [
                    str(v)[: settings.max_cell_chars].strip() if v is not None else ""
                    for v in row
                ]
                if not any(cells):
                    continue
                if headers is None:
                    # Az első nem üres sor a fejléc
                    headers = cells
                    continue
                filled = [(i, c) for i, c in enumerate(cells) if c]
                # Kevés kitöltött cella (pl. megjegyzés-sor): fejléc-címke nélkül,
                # különben félrevezető párosítás születne
                if len(filled) <= 2:
                    yield " — ".join(c for _, c in filled)
                    continue
                pairs = []
                for i, c in filled:
                    header = headers[i] if i < len(headers) and headers[i] else f"{i + 1}. oszlop"
                    pairs.append(f"{header}: {c}")
                yield " | ".join(pairs)
    finally:
        wb.close()


LOADERS = {
    ".txt": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
}


def extract_text(filename: str, data: bytes) -> Iterator[str]:
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    loader = LOADERS.get(ext)
    if loader is None:
        raise ValueError(f"Nem támogatott fájltípus: {ext or filename}")
    return loader(data)
